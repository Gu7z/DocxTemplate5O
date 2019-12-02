from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import Inches
from PIL import Image
import pyqrcode
import time
import os
import re 

def tamanho(string):
    if( len(string) > 13 ):
            return True
    else:
        print('----------------')
        print('O selo não é válido.')
        print('Os selos devem ter pelo menos 14 caracteres.')
        time.sleep(3)
        return False

def procurar(string): 
    regex = re.compile('-') 
    if(regex.search(string) == None):
        if tamanho(string):
            return True
        else: 
            return False
    else: 
        global selo
        string = string.replace("-", "")
        string = string.replace(" ", "")
        selo = string
        if tamanho(string):
            return True
        else: 
            return False
try:
    img = Image.open('./qrcode.bmp')
    new_img = img.resize(( 256,256 ))
    new_img.save( './qrcode.png', 'png' )
except:
    print('O selo não está na pasta')

selo = input('Digite o selo: ')

if procurar(selo):
    selo = selo.upper()
    selo = selo[:len(selo)-6] + '-' + selo[len(selo)-6:]
    selo = selo[:len(selo)-3] + '-' + selo[len(selo)-3:]

    document = Document()
    styles = document.styles

    section = document.sections[0]
    section.page_height = Cm(2)
    section.page_width = Cm(5.5)
    section.top_margin = Cm(0.1)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0.1)
    section.right_margin = Cm(0)


    f1 = document.add_paragraph('')
    run1 = f1.add_run('Selo Digital de Autenticidade:       ')
    style = styles.add_style('Frase1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(6)
    f1.style = document.styles['Frase1']
    try: 
        run1.add_picture('qrcode.png', width=Cm(1.5))
    except:
        print('Seu documento estará sem selo')

    f2 = document.add_paragraph()
    run2 = f2.add_run(selo)
    style2 = styles.add_style('Frase2', WD_STYLE_TYPE.PARAGRAPH)
    font2 = style2.font
    font2.name = 'Times New Roman'
    font2.size = Pt(10)
    f2.style = document.styles['Frase2']

    f3 = document.add_paragraph()
    run3 = f3.add_run('http://www.tjms.jus.br/corregedoria/selos')
    style3 = styles.add_style('Frase3', WD_STYLE_TYPE.PARAGRAPH)
    font3 = style3.font
    font3.name = 'Times New Roman'
    font3.size = Pt(8)
    f3.style = document.styles['Frase3']

    run1.bold = True
    run2.bold = True
    run3.bold = True

    pf1 = f1.paragraph_format
    pf2 = f2 .paragraph_format

    pf1.left_indent = Cm(0.75)
    pf2.left_indent = Cm(0.25)

    document.save('{}.docx'.format(selo))
